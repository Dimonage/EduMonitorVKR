# -*- coding: utf-8 -*-
"""
Модуль utils.py
Описание: Вспомогательные функции для анализа данных в системе мониторинга вузов.
Содержит функции для проверки качества данных, визуализации признаков, анализа важности признаков и экспорта в Word.
Используется в приложении Streamlit для расширения функциональности.
"""

import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import streamlit as st
import os
from docx import Document
from docx.shared import Inches
from scipy import stats

# Настройка стиля графиков
sns.set(style="whitegrid")
plt.rcParams["figure.figsize"] = (10, 6)

# Создание директории для графиков
os.makedirs("plots", exist_ok=True)

# Список для хранения путей к сохранённым графикам
saved_plots = []
# Список для хранения читаемых описаний графиков
plot_descriptions = []

def clear_saved_plots():
    """
    Очистка списка сохранённых графиков и удаление файлов из директории plots
    """
    try:
        for file in os.listdir("plots"):
            os.remove(os.path.join("plots", file))
        saved_plots.clear()
        plot_descriptions.clear()
        st.write("Сохранённые графики и описания очищены")
        return True
    except Exception as e:
        st.error(f"Ошибка при очистке графиков: {e}")
        return False

def check_data_quality(df):
    """
    Проверка качества данных: пропуски, статистика, типы данных
    
    Параметры:
    df (pandas.DataFrame): Входной датасет
    
    Возвращает:
    dict: Словарь с информацией о пропусках, статистике и типах данных
    """
    st.subheader("Анализ качества данных")
    
    # Проверка пропусков
    missing_values = df.isna().sum()
    missing_summary = pd.DataFrame({
        'Столбец': missing_values.index,
        'Количество пропусков': missing_values.values,
        'Процент пропусков': (missing_values / len(df) * 100).round(2)
    })
    st.write("Пропуски в данных:")
    st.dataframe(missing_summary)
    
    # Основная статистика
    stats = df.describe().round(2)
    st.write("Основная статистика по числовым столбцам:")
    st.dataframe(stats)
    
    # Типы данных
    dtypes_summary = pd.DataFrame({
        'Столбец': df.columns,
        'Тип данных': df.dtypes.values
    })
    st.write("Типы данных столбцов:")
    st.dataframe(dtypes_summary)
    
    result = {
        'missing': missing_summary,
        'stats': stats,
        'dtypes': dtypes_summary
    }
    return result

def plot_feature_distribution(df, feature, save_path=None):
    """
    Построение гистограммы распределения признака с ядерной оценкой плотности
    
    Параметры:
    df (pandas.DataFrame): Входной датасет
    feature (str): Название столбца для анализа
    save_path (str, optional): Путь для сохранения графика
    
    Возвращает:
    str: Путь к сохранённому графику или None при ошибке
    """
    try:
        if feature not in df.columns:
            st.error(f"Признак '{feature}' отсутствует в данных")
            return None
        if not np.issubdtype(df[feature].dtype, np.number):
            st.error("Признак должен быть числовым")
            return None
        if df[feature].isna().all():
            st.error("Признак содержит только пропуски")
            return None
        
        st.subheader(f"Распределение признака '{feature}'")
        fig, ax = plt.subplots()
        sns.histplot(df[feature], kde=True, ax=ax, color='skyblue')
        ax.set_title(f"Распределение признака '{feature}'")
        ax.set_xlabel(feature)
        ax.set_ylabel("Частота")
        st.pyplot(fig)
        
        # Автоматическое сохранение графика
        if save_path is None:
            save_path = f"plots/feature_distribution_{feature.replace(' ', '_')}_{len(saved_plots)}.png"
        os.makedirs(os.path.dirname(save_path), exist_ok=True)
        fig.savefig(save_path, bbox_inches='tight')
        saved_plots.append(save_path)
        # Добавление читаемого описания
        description = f"Распределение признака: {feature}"
        plot_descriptions.append(description)
        st.write(f"График сохранён как {save_path}")
        
        plt.close(fig)
        return save_path
    except Exception as e:
        st.error(f"Ошибка при построении гистограммы: {e}")
        return None

def plot_feature_importance(model, feature_names, save_path=None):
    """
    Построение графика важности признаков для модели регрессии
    
    Параметры:
    model: Обученная модель с атрибутом feature_importances_
    feature_names (list): Список названий признаков
    save_path (str, optional): Путь для сохранения графика
    
    Возвращает:
    str: Путь к сохранённому графику или None при ошибке
    """
    try:
        if not hasattr(model, 'feature_importances_'):
            st.error("Модель не поддерживает анализ важности признаков")
            return None
        if len(feature_names) != len(model.feature_importances_):
            st.error("Количество признаков не соответствует модели")
            return None
        
        st.subheader("Важность признаков модели")
        fig, ax = plt.subplots(figsize=(10, 8))
        sns.barplot(x=model.feature_importances_, y=feature_names, ax=ax, palette='viridis')
        ax.set_title("Важность признаков в модели")
        ax.set_xlabel("Важность")
        ax.set_ylabel("Признак")
        st.pyplot(fig)
        
        # Автоматическое сохранение графика
        if save_path is None:
            save_path = f"plots/feature_importance_{len(saved_plots)}.png"
        os.makedirs(os.path.dirname(save_path), exist_ok=True)
        fig.savefig(save_path, bbox_inches='tight')
        saved_plots.append(save_path)
        # Добавление читаемого описания
        description = "Важность признаков в модели регрессии"
        plot_descriptions.append(description)
        st.write(f"График сохранён как {save_path}")
        
        plt.close(fig)
        return save_path
    except Exception as e:
        st.error(f"Ошибка при построении графика важности: {e}")
        return None

def save_dataframe(df, filename="data.csv"):
    """
    Сохранение датасета в файл
    
    Параметры:
    df (pandas.DataFrame): Датасет для сохранения
    filename (str): Имя файла
    
    Возвращает:
    bool: True, если сохранение успешно
    """
    try:
        df.to_csv(filename, index=False)
        st.write(f"Датасет сохранён как {filename}")
        return True
    except Exception as e:
        st.error(f"Ошибка при сохранении датасета: {e}")
        return False

def load_dataframe(filename="data.csv"):
    """
    Загрузка датасета из файла
    
    Параметры:
    filename (str): Имя файла
    
    Возвращает:
    pandas.DataFrame или None: Загруженный датасет или None при ошибке
    """
    try:
        if not os.path.exists(filename):
            st.error(f"Файл {filename} не найден")
            return None
        df = pd.read_csv(filename)
        st.write(f"Датасет загружен из {filename}")
        return df
    except Exception as e:
        st.error(f"Ошибка при загрузке датасета: {e}")
        return None

def detect_outliers(df, column, iqr_factor=1.5):
    """
    Обнаружение выбросов с помощью метода межквартильного размаха
    
    Параметры:
    df (pandas.DataFrame): Входной датасет
    column (str): Столбец для анализа
    iqr_factor (float): Множитель для IQR (по умолчанию 1.5)
    
    Возвращает:
    pandas.Series: Булева маска, где True — выбросы
    """
    try:
        if column not in df.columns:
            st.error(f"Столбец '{column}' отсутствует")
            return None
        if not np.issubdtype(df[column].dtype, np.number):
            st.error("Столбец должен быть числовым")
            return None
        
        q1 = df[column].quantile(0.25)
        q3 = df[column].quantile(0.75)
        iqr = q3 - q1
        lower_bound = q1 - iqr_factor * iqr
        upper_bound = q3 + iqr_factor * iqr
        outliers = (df[column] < lower_bound) | (df[column] > upper_bound)
        
        st.write(f"Обнаружено {outliers.sum()} выбросов в столбце '{column}'")
        return outliers
    except Exception as e:
        st.error(f"Ошибка при обнаружении выбросов: {e}")
        return None

def summarize_features(df, max_features=10):
    """
    Краткий анализ признаков: среднее, дисперсия, уникальные значения
    
    Параметры:
    df (pandas.DataFrame): Входной датасет
    max_features (int): Максимальное количество признаков для анализа
    
    Возвращает:
    pandas.DataFrame: Таблица с характеристиками признаков
    """
    try:
        summary = []
        for col in df.columns[:max_features]:
            mean_val = df[col].mean() if np.issubdtype(df[col].dtype, np.number) else None
            var_val = df[col].var() if np.issubdtype(df[col].dtype, np.number) else None
            unique_count = df[col].nunique()
            summary.append({
                'Признак': col,
                'Среднее': round(mean_val, 2) if mean_val is not None else '-',
                'Дисперсия': round(var_val, 2) if var_val is not None else '-',
                'Уникальных значений': unique_count
            })
        summary_df = pd.DataFrame(summary)
        st.write("Характеристики признаков:")
        st.dataframe(summary_df)
        return summary_df
    except Exception as e:
        st.error(f"Ошибка при анализе признаков: {e}")
        return None
    
def check_duplicates(df):
    """
    Проверка датасета на наличие дубликатов.
    
    Параметры:
    df (pandas.DataFrame): Входной датасет
    
    Возвращает:
    dict: Информация о дубликатах (количество и строки)
    """
    try:
        duplicates = df.duplicated().sum()
        duplicate_rows = df[df.duplicated()].head()
        result = {
            'count': duplicates,
            'rows': duplicate_rows
        }
        st.write(f"Найдено дубликатов: {duplicates}")
        if duplicates > 0:
            st.write("Примеры дублированных строк:")
            st.dataframe(duplicate_rows)
        return result
    except Exception as e:
        st.error(f"Ошибка при проверке дубликатов: {e}")
        return None

def remove_duplicates(df):
    """
    Удаление дубликатов из датасета.
    
    Параметры:
    df (pandas.DataFrame): Входной датасет
    
    Возвращает:
    pandas.DataFrame: Датасет без дубликатов
    """
    try:
        initial_shape = df.shape
        df_clean = df.drop_duplicates()
        st.write(f"Удалено дубликатов: {initial_shape[0] - df_clean.shape[0]}")
        st.write(f"Новый размер датасета: {df_clean.shape}")
        return df_clean
    except Exception as e:
        st.error(f"Ошибка при удалении дубликатов: {e}")
        return df
    
def filter_data(df, column, value):
    """
    Фильтрация датасета по значению в указанном столбце.
    
    Параметры:
    df (pandas.DataFrame): Входной датасет
    column (str): Столбец для фильтрации
    value: Значение для фильтрации
    
    Возвращает:
    pandas.DataFrame: Отфильтрованный датасет
    """
    try:
        if column not in df.columns:
            st.error(f"Столбец '{column}' отсутствует")
            return df
        filtered_df = df[df[column] == value]
        st.write(f"Отфильтровано {filtered_df.shape[0]} строк по {column} = {value}")
        return filtered_df
    except Exception as e:
        st.error(f"Ошибка при фильтрации: {e}")
        return df
    
def export_filtered_data(df, filename="filtered_data_export.csv"):
    """
    Экспорт отфильтрованного датасета в CSV.
    
    Параметры:
    df (pandas.DataFrame): Датасет для экспорта
    filename (str): Имя файла
    
    Возвращает:
    bool: True, если экспорт успешен
    """
    try:
        df.to_csv(filename, index=False)
        st.write(f"Отфильтрованный датасет сохранён как {filename}")
        return True
    except Exception as e:
        st.error(f"Ошибка при экспорте: {e}")
        return False

def calculate_selected_correlations(df, columns):
    """
    Расчёт корреляций для выбранных столбцов.
    
    Параметры:
    df (pandas.DataFrame): Входной датасет
    columns (list): Список столбцов
    
    Возвращает:
    pandas.DataFrame: Матрица корреляций
    """
    try:
        numeric_cols = [col for col in columns if np.issubdtype(df[col].dtype, np.number)]
        if not numeric_cols:
            st.error("Выбранные столбцы не содержат числовых данных")
            return None
        corr_matrix = df[numeric_cols].corr()
        st.write("Матрица корреляций для выбранных столбцов:")
        st.dataframe(corr_matrix)
        return corr_matrix
    except Exception as e:
        st.error(f"Ошибка при расчёте корреляций: {e}")
        return None
    
def export_descriptive_stats(stats_dict, column, filename="descriptive_stats.xlsx"):
    """
    Экспорт описательных статистик в Excel.
    
    Параметры:
    stats_dict (dict): Словарь со статистиками
    column (str): Название столбца
    filename (str): Имя файла
    
    Возвращает:
    bool: True, если экспорт успешен
    """
    try:
        if stats_dict is None:
            st.error("Нет данных для экспорта")
            return False
        stats_df = pd.DataFrame.from_dict(stats_dict, orient='index', columns=['Значение'])
        stats_df['Значение'] = stats_df['Значение'].round(2)
        stats_df.to_excel(filename, sheet_name=f"Stats_{column}", engine='xlsxwriter')
        st.write(f"Статистики сохранены в {filename}")
        return True
    except Exception as e:
        st.error(f"Ошибка при экспорте статистик: {e}")
        return False
    
def calculate_descriptive_stats(df, column):
    """
    Расчёт описательных статистик для выбранного столбца.
    
    Параметры:
    df (pandas.DataFrame): Входной датасет
    column (str): Столбец для анализа
    
    Возвращает:
    dict: Словарь с описательными статистиками
    """
    try:
        if column not in df.columns:
            st.error(f"Столбец '{column}' отсутствует")
            return None
        if not np.issubdtype(df[column].dtype, np.number):
            st.error("Столбец должен быть числовым")
            return None
        
        stats_dict = {
            'Среднее': df[column].mean(),
            'Медиана': df[column].median(),
            'Мода': stats.mode(df[column], keepdims=True)[0][0],
            'Стд. отклонение': df[column].std(),
            'Минимум': df[column].min(),
            'Максимум': df[column].max()
        }
        stats_df = pd.DataFrame.from_dict(stats_dict, orient='index', columns=['Значение'])
        stats_df['Значение'] = stats_df['Значение'].round(2)
        st.write(f"Описательные статистики для столбца '{column}':")
        st.dataframe(stats_df)
        return stats_dict
    except Exception as e:
        st.error(f"Ошибка при расчёте статистик: {e}")
        return None
    
def export_to_word(output_file="edu_monitor_report.docx", standard_text=None):
    """
    Экспорт результатов анализа в документ Word с текстом и графиками
    
    Параметры:
    output_file (str): Имя выходного файла Word
    standard_text (str): Стандартизированный текст для включения в документ
    
    Возвращает:
    bool: True, если экспорт успешен
    """
    try:
        doc = Document()
        
        # Добавление заголовка
        doc.add_heading("Отчёт по анализу данных мониторинга вузов", 0)
        
        # Добавление стандартизированного текста
        if standard_text:
            doc.add_paragraph(standard_text)
        else:
            doc.add_paragraph(
                "Данный отчёт содержит результаты интеллектуального анализа данных мониторинга вузов, "
                "включая анализ качества данных, распределение признаков, важность признаков, "
                "результаты кластеризации и корреляционный анализ. Графики ниже иллюстрируют основные выводы."
            )
        
        # Добавление графиков
        doc.add_heading("Визуализации", level=1)
        for plot_path, description in zip(saved_plots, plot_descriptions):
            if os.path.exists(plot_path):
                doc.add_paragraph(f"График: {description}")
                doc.add_picture(plot_path, width=Inches(6))
            else:
                st.warning(f"Файл {plot_path} не найден и не включён в отчёт")
        
        # Сохранение документа
        doc.save(output_file)
        st.write(f"Отчёт сохранён как {output_file}")
        return True
    except Exception as e:
        st.error(f"Ошибка при создании документа Word: {e}")
        return False